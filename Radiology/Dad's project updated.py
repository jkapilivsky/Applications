# values
'''
name
Date of Exam
DOB
medical record:
Referral DR.
'''
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import xlrd
import datetime
import os

excel_file = 'Patient info.xlsx'
tab_name = 'Sheet1'

book = xlrd.open_workbook(excel_file)
sheet = book.sheet_by_name(tab_name)
data = [[sheet.cell_value(r, c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]

for reports in range(1, len(data)):
    dob_text = data[reports][1]
    dob_sec = (dob_text - 25569) * 86400.0
    dob_time = datetime.datetime.utcfromtimestamp(dob_sec)
    s = str(dob_time)
    DOB = s[:-9]
    new_DOB = DOB.split(sep='-')
    final_DOB = new_DOB[1] + '/' + new_DOB[2] + '/' + new_DOB[0][2:4]


    exam_text = data[reports][2]
    exam_sec = (exam_text - 25569) * 86400.0
    exam_time = datetime.datetime.utcfromtimestamp(exam_sec)
    s2 = str(exam_time)
    date_of_exam = s2[:-9]
    new_date_of_exam = date_of_exam.split(sep='-')
    final_date_of_exam = new_date_of_exam[1] + '/' + new_date_of_exam[2] + '/' + new_date_of_exam[0][2:4]

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    MRI_info = document.add_paragraph()
    MRI_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    MRI_info.add_run().add_picture('Assets/MRI logo.png', width=Inches(1.25))
    MRI_info.add_run().add_break()
    MRI_info.add_run('McAllen MRI Center').bold = True
    MRI_info.add_run().add_break()
    MRI_info.add_run('320 North McColl Rd. E')
    MRI_info.add_run().add_break()
    MRI_info.add_run('McAllen, TX 78501')
    MRI_info.add_run().add_break()
    MRI_info.add_run('Phone: (956)-687-9636  Fax: (956)-687-9743')

##############################################################################
    patient_info = document.add_paragraph('-' * 108)
    patient_info.paragraph_format.space_after = Pt(0)
    patient_info.add_run().add_break()
    patient_info.add_run('NAME: ').bold = True
    patient_info.add_run(data[reports][0])


    #patient_info.add_run('  '*(30-len('NAME: ' + data[reports][0])))
    if len(data[reports][0]) >= 10:
        patient_info.add_run().add_tab()
        patient_info.add_run().add_tab()
        patient_info.add_run().add_tab()
    else:
        patient_info.add_run().add_tab()
        patient_info.add_run().add_tab()
        patient_info.add_run().add_tab()
        patient_info.add_run().add_tab()

    patient_info.add_run('Date of Exam: ').bold = True
    patient_info.add_run(final_date_of_exam)
    patient_info.add_run().add_break()


    patient_info.add_run('DOB: ').bold = True
    patient_info.add_run(final_DOB)

    #patient_info.add_run('  '*(30-len('DOB: ' + DOB)))
    patient_info.add_run().add_tab()
    patient_info.add_run().add_tab()
    patient_info.add_run().add_tab()
    patient_info.add_run().add_tab()


    patient_info.add_run('Medical Record: ').bold = True
    # Allows strings and floats as data requests
    if type(data[reports][4]) is str:
        patient_info.add_run(data[reports][4])
    elif type(data[reports][4]) is float:
        patient_info.add_run(str(int(data[reports][4])))

    patient_info.add_run().add_break()


    patient_info.add_run('Referral Dr: ').bold = True
    patient_info.add_run(data[reports][3])
    patient_info.add_run('  '*(30-len('Referral Dr: ' + data[reports][3])))
    # patient_info.add_run().add_tab()
    # patient_info.add_run().add_tab()

    patient_info.add_run().add_break()

    patient_info.add_run('Procedure: ').bold = True
    patient_info.add_run(data[reports][5])
    patient_info.add_run().add_break()
    patient_info.add_run().add_break()

    patient_info.add_run('History: ').bold = True
    patient_info.space_after = Pt(0)
    patient_info.add_run().add_break()
    patient_info.add_run().add_break()

    patient_info.add_run('-' * 108)


    ##############################################################################
    patient_info.add_run().add_break()
    signature = document.add_paragraph('Thank you for your referral')
    signature.add_run().add_break()
    signature.add_run().add_picture('Assets/Signature.jpg')
    signature.add_run().add_break()
    signature.add_run('Electronically signed by:')
    signature.add_run().add_break()
    signature.add_run('Allan Kapilivsky M.D.')
    signature.add_run().add_break()
    signature.add_run('Reviewed and dictated the same day')

    # Saving document
    folder = 'Reports/' + str(datetime.date.today()) + '/'
    if not os.path.exists(folder):
        os.makedirs(folder)

    document.save(folder + data[reports][0] + " " + date_of_exam + '.docx')
print('Complete')


# TODO - wants new version for each doctor


'''V1.1 Done'''
# TODO - put procedure and history ablve 2nd dotted line
# TODO - times new roman 12pt
# TODO - remove double space before signature!
# TODO - make signature smaller
# TODO - Medical Record needs to accept numbers and letters
# TODO - 6/14/67 new format
# TODO - fixed tab issues to be depended on patient length


